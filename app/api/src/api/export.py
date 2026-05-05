"""Export API - Generate PDF, Word, and Excel exports of RFP responses."""

from datetime import datetime
from io import BytesIO
from typing import Annotated
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

from src.domain.enums import AnswerStatus
from src.domain.schemas import UserResponse
from src.infrastructure.database.connection import get_db
from src.api.dependencies import get_current_user
from src.infrastructure.database.repositories import (
    SQLAlchemyAnswerRepository,
    SQLAlchemyProjectRepository,
    SQLAlchemyQuestionRepository,
)

router = APIRouter(prefix="/export", tags=["export"])


@router.get("/project/{project_id}/pdf")
async def export_project_pdf(
    project_id: UUID,
    session: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[UserResponse, Depends(get_current_user)],
):
    """Export all questions and answers for a project as PDF.
    
    Requires authentication and project ownership (tenant isolation).
    """
    # Verify project access with tenant isolation
    project_repo = SQLAlchemyProjectRepository(session)
    project = await project_repo.get_by_id(project_id, tenant_id=current_user.tenant_id)
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found or access denied",
        )
    
    # Get questions and answers
    question_repo = SQLAlchemyQuestionRepository(session)
    answer_repo = SQLAlchemyAnswerRepository(session)

    # SECURITY: tenant_id obrigatório para isolamento
    questions = await question_repo.get_by_project(project_id, tenant_id=current_user.tenant_id)

    # Build QA pairs
    qa_pairs = []
    for question in questions:
        # SECURITY: tenant_id obrigatório para isolamento
        answer = await answer_repo.get_by_question(question.id, tenant_id=current_user.tenant_id)
        qa_pairs.append({
            "question": question,
            "answer": answer,
        })

    # Generate HTML content
    html_content = generate_html_export(project, qa_pairs)

    # Convert to PDF using WeasyPrint
    try:
        from weasyprint import HTML
        pdf_buffer = BytesIO()
        HTML(string=html_content).write_pdf(pdf_buffer)
        pdf_buffer.seek(0)

        filename = f"rfp_response_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
        
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate PDF: {str(e)}",
        )


@router.get("/project/{project_id}/word")
async def export_project_word(
    project_id: UUID,
    session: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[UserResponse, Depends(get_current_user)],
):
    """Export all questions and answers for a project as Word document.
    
    Requires authentication and project ownership (tenant isolation).
    """
    # Verify project access with tenant isolation
    project_repo = SQLAlchemyProjectRepository(session)
    project = await project_repo.get_by_id(project_id, tenant_id=current_user.tenant_id)
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found or access denied",
        )
    
    # Get questions and answers
    question_repo = SQLAlchemyQuestionRepository(session)
    answer_repo = SQLAlchemyAnswerRepository(session)

    # SECURITY: tenant_id obrigatório para isolamento
    questions = await question_repo.get_by_project(project_id, tenant_id=current_user.tenant_id)

    # Build QA pairs
    qa_pairs = []
    for question in questions:
        # SECURITY: tenant_id obrigatório para isolamento
        answer = await answer_repo.get_by_question(question.id, tenant_id=current_user.tenant_id)
        qa_pairs.append({
            "question": question,
            "answer": answer,
        })

    # Generate Word document
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f"RFP Response: {project.name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total Questions: {len(questions)}")
        doc.add_paragraph()
        
        # Questions and Answers
        for i, qa in enumerate(qa_pairs, 1):
            question = qa["question"]
            answer = qa["answer"]
            
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{i}: {question.question_text}")
            q_run.bold = True
            q_run.font.size = Pt(11)
            
            # Category
            doc.add_paragraph(f"Category: {question.category.value}", style='Intense Quote')
            
            # Answer
            if answer:
                answer_text = answer.final_text or answer.suggested_text or "No answer provided"
                a_para = doc.add_paragraph()
                a_run = a_para.add_run("Answer: ")
                a_run.bold = True
                a_para.add_run(answer_text)
                
                # Status
                status_para = doc.add_paragraph()
                status_para.add_run(f"Status: {answer.status.value}")
                if answer.confidence:
                    status_para.add_run(f" | Confidence: {answer.confidence:.0%}")
            else:
                doc.add_paragraph("Answer: Pending generation", style='Quote')
            
            doc.add_paragraph()  # Spacing
        
        # Save to buffer
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)
        
        filename = f"rfp_response_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            word_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate Word document: {str(e)}",
        )


@router.get("/project/{project_id}/excel")
async def export_project_excel(
    project_id: UUID,
    session: Annotated[AsyncSession, Depends(get_db)],
    current_user: Annotated[UserResponse, Depends(get_current_user)],
):
    """Export all questions and answers for a project as Excel spreadsheet.
    
    Requires authentication and project ownership (tenant isolation).
    """
    # Verify project access with tenant isolation
    project_repo = SQLAlchemyProjectRepository(session)
    project = await project_repo.get_by_id(project_id, tenant_id=current_user.tenant_id)
    if not project:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Project not found or access denied",
        )
    
    # Get questions and answers
    question_repo = SQLAlchemyQuestionRepository(session)
    answer_repo = SQLAlchemyAnswerRepository(session)

    # SECURITY: tenant_id obrigatório para isolamento
    questions = await question_repo.get_by_project(project_id, tenant_id=current_user.tenant_id)

    # Build data for Excel
    data = []
    for i, question in enumerate(questions, 1):
        # SECURITY: tenant_id obrigatório para isolamento
        answer = await answer_repo.get_by_question(question.id, tenant_id=current_user.tenant_id)

        data.append({
            "#": i,
            "Question": question.question_text,
            "Category": question.category.value,
            "Priority": question.priority,
            "Answer": answer.final_text or answer.suggested_text if answer else "",
            "Status": answer.status.value if answer else "pending",
            "Confidence": f"{answer.confidence:.0%}" if answer and answer.confidence else "",
            "Approved": "Yes" if answer and answer.status == AnswerStatus.APPROVED else "No",
        })
    
    # Generate Excel
    try:
        import pandas as pd
        
        df = pd.DataFrame(data)
        
        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='RFP Responses', index=False)
            
            # Get workbook and worksheet for formatting
            workbook = writer.book
            worksheet = writer.sheets['RFP Responses']
            
            # Adjust column widths
            worksheet.column_dimensions['A'].width = 5   # #
            worksheet.column_dimensions['B'].width = 50  # Question
            worksheet.column_dimensions['C'].width = 15  # Category
            worksheet.column_dimensions['D'].width = 10  # Priority
            worksheet.column_dimensions['E'].width = 50  # Answer
            worksheet.column_dimensions['F'].width = 12  # Status
            worksheet.column_dimensions['G'].width = 12  # Confidence
            worksheet.column_dimensions['H'].width = 12  # Approved
        
        excel_buffer.seek(0)
        
        filename = f"rfp_response_{project.name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        return StreamingResponse(
            excel_buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate Excel: {str(e)}",
        )


def generate_html_export(project, qa_pairs) -> str:
    """Generate HTML content for PDF export."""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>RFP Response - {project.name}</title>
        <style>
            body {{
                font-family: Arial, sans-serif;
                margin: 40px;
                line-height: 1.6;
            }}
            h1 {{
                color: #333;
                text-align: center;
                border-bottom: 2px solid #007bff;
                padding-bottom: 10px;
            }}
            .metadata {{
                text-align: center;
                color: #666;
                margin-bottom: 30px;
            }}
            .qa-item {{
                margin-bottom: 25px;
                page-break-inside: avoid;
            }}
            .question {{
                font-weight: bold;
                color: #007bff;
                margin-bottom: 5px;
            }}
            .category {{
                font-size: 0.9em;
                color: #666;
                font-style: italic;
                margin-bottom: 10px;
            }}
            .answer {{
                margin-left: 20px;
                padding: 10px;
                background-color: #f8f9fa;
                border-left: 3px solid #28a745;
            }}
            .pending {{
                border-left-color: #ffc107;
                background-color: #fff3cd;
            }}
            .status {{
                font-size: 0.85em;
                color: #28a745;
                margin-top: 5px;
            }}
            .approved {{
                color: #28a745;
                font-weight: bold;
            }}
            .page-break {{
                page-break-after: always;
            }}
        </style>
    </head>
    <body>
        <h1>RFP Response: {project.name}</h1>
        <div class="metadata">
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            <p>Total Questions: {len(qa_pairs)}</p>
        </div>
    """
    
    for i, qa in enumerate(qa_pairs, 1):
        question = qa["question"]
        answer = qa["answer"]
        
        answer_text = ""
        status_class = "pending"
        status_text = "Pending"
        
        if answer:
            answer_text = answer.final_text or answer.suggested_text or ""
            status_class = "" if answer.status == AnswerStatus.APPROVED else "pending"
            status_text = f"Status: {answer.status.value}"
            if answer.confidence:
                status_text += f" | Confidence: {answer.confidence:.0%}"
        
        html += f"""
        <div class="qa-item">
            <div class="question">Q{i}: {question.question_text}</div>
            <div class="category">Category: {question.category.value} | Priority: {question.priority}</div>
            <div class="answer {status_class}">
                {answer_text or "<em>No answer provided</em>"}
            </div>
            <div class="status {"approved" if answer and answer.status == AnswerStatus.APPROVED else ""}">{status_text}</div>
        </div>
        """
    
    html += """
    </body>
    </html>
    """
    
    return html
