import io
import re
from pathlib import Path
from typing import Any

import pdfplumber
import pandas as pd
from docx import Document as DocxDocument

from src.core.logging_config import get_logger
from src.domain.enums import RFPType
from src.domain.exceptions import DocumentProcessingException

logger = get_logger(__name__)


def _detect_encoding(file_path: str) -> tuple[str, bytes]:
    """Detect file encoding and return (encoding, raw_bytes).
    
    Prioritizes UTF-8 for proper handling of Portuguese characters.
    Tries multiple strategies:
    1. UTF-8 BOM check
    2. UTF-8 strict decode
    3. UTF-8 with replacement
    4. Fallback encodings only if UTF-8 fails
    """
    with open(file_path, 'rb') as f:
        raw_bytes = f.read()
    
    # Check for UTF-8 BOM
    if raw_bytes.startswith(b'\xef\xbb\xbf'):
        return 'utf-8-sig', raw_bytes
    
    # Try UTF-8 first (most common for Portuguese text)
    try:
        # Test if it's valid UTF-8
        decoded = raw_bytes.decode('utf-8')
        # Additional check: if it contains Portuguese characters, it's likely UTF-8
        if any(char in decoded for char in 'çãõáéíóúâêîôûÇÃÕÁÉÍÓÚÂÊÎÔÛ'):
            return 'utf-8', raw_bytes
        return 'utf-8', raw_bytes
    except UnicodeDecodeError:
        pass
    
    # Try UTF-8 with BOM signature
    try:
        raw_bytes.decode('utf-8-sig')
        return 'utf-8-sig', raw_bytes
    except UnicodeDecodeError:
        pass
    
    # Only try legacy encodings as fallback
    fallback_encodings = ['cp1252', 'latin-1', 'iso-8859-1']
    
    for encoding in fallback_encodings:
        try:
            raw_bytes.decode(encoding)
            return encoding, raw_bytes
        except UnicodeDecodeError:
            continue
    
    # Ultimate fallback to UTF-8 with replacement
    return 'utf-8', raw_bytes


class DocumentParser:
    """Parse various document formats into text."""

    @staticmethod
    def parse_pdf(file_path: str) -> dict[str, Any]:
        """Extract text and metadata from PDF."""
        try:
            text_parts = []
            metadata = {"page_count": 0, "pages": []}

            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        metadata["pages"].append({
                            "page_number": i + 1,
                            "char_count": len(page_text),
                        })

            return {
                "text": "\n\n".join(text_parts),
                "metadata": metadata,
            }
        except Exception as e:
            logger.error("PDF parsing failed", file_path=file_path, error=str(e))
            raise DocumentProcessingException(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> dict[str, Any]:
        """Extract text and metadata from DOCX."""
        try:
            doc = DocxDocument(file_path)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))

            full_text = "\n\n".join(text_parts)
            if tables_text:
                full_text += "\n\n--- Tables ---\n\n"
                full_text += "\n\n".join(tables_text)

            return {
                "text": full_text,
                "metadata": {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                },
            }
        except Exception as e:
            logger.error("DOCX parsing failed", file_path=file_path, error=str(e))
            raise DocumentProcessingException(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def parse_xlsx(file_path: str) -> dict[str, Any]:
        """Extract text and metadata from XLSX."""
        try:
            df = pd.read_excel(file_path, sheet_name=None)

            all_sheets_text = []
            for sheet_name, sheet_df in df.items():
                sheet_text = f"Sheet: {sheet_name}\n"
                sheet_text += sheet_df.to_string(index=False)
                all_sheets_text.append(sheet_text)

            return {
                "text": "\n\n---\n\n".join(all_sheets_text),
                "metadata": {
                    "sheet_count": len(df),
                    "sheet_names": list(df.keys()),
                },
            }
        except Exception as e:
            logger.error("XLSX parsing failed", file_path=file_path, error=str(e))
            raise DocumentProcessingException(f"Failed to parse XLSX: {str(e)}")

    @staticmethod
    def parse_txt(file_path: str) -> dict[str, Any]:
        """Extract text from plain text file with automatic encoding detection."""
        try:
            # Detect encoding automatically
            detected_encoding, raw_bytes = _detect_encoding(file_path)
            
            # Check if file is empty before decoding
            if len(raw_bytes) == 0:
                raise DocumentProcessingException("File is empty (0 bytes)")
            
            # Decode with detected encoding, using replacement for invalid chars
            text = raw_bytes.decode(detected_encoding, errors='replace')
            
            # Clean up common encoding artifacts
            text = text.replace('\x00', '')  # Remove null bytes
            text = text.replace('\r\n', '\n')  # Normalize line endings
            text = text.replace('\r', '\n')
            
            # Check if text is empty after cleaning
            if len(text.strip()) == 0:
                raise DocumentProcessingException("File contains no readable text content")
            
            logger.info(
                "TXT parsed successfully",
                file_path=file_path,
                encoding=detected_encoding,
                char_count=len(text),
            )

            return {
                "text": text,
                "metadata": {
                    "char_count": len(text),
                    "line_count": len(text.splitlines()),
                    "encoding": detected_encoding,
                },
            }
        except DocumentProcessingException:
            # Re-raise our own exceptions
            raise
        except Exception as e:
            logger.error(
                "TXT parsing failed",
                file_path=file_path,
                error=str(e),
                error_type=type(e).__name__,
            )
            raise DocumentProcessingException(f"Failed to parse TXT: {str(e)}")

    @classmethod
    def parse(cls, file_path: str, doc_type: str) -> dict[str, Any]:
        """Parse document based on type."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return cls.parse_pdf(file_path)
        elif ext == ".docx":
            return cls.parse_docx(file_path)
        elif ext in [".xlsx", ".xls"]:
            return cls.parse_xlsx(file_path)
        elif ext == ".txt":
            return cls.parse_txt(file_path)
        else:
            raise DocumentProcessingException(f"Unsupported file type: {ext}")


class TextChunker:
    """Chunk text into smaller pieces with overlap."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk(self, text: str) -> list[dict[str, Any]]:
        """Split text into overlapping chunks."""
        # Simple sentence-aware chunking with Unicode support
        sentences = re.split(r'(?<=[.!?])\s+', text, flags=re.UNICODE)

        chunks = []
        current_chunk = []
        current_size = 0
        chunk_index = 0

        for sentence in sentences:
            sentence_size = len(sentence)

            if current_size + sentence_size > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": chunk_index,
                    "char_count": len(chunk_text),
                })
                chunk_index += 1

                # Start new chunk with overlap
                overlap_size = 0
                overlap_sentences = []
                for s in reversed(current_chunk):
                    if overlap_size + len(s) <= self.chunk_overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s) + 1
                    else:
                        break

                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s) for s in current_chunk) + len(current_chunk) - 1
            else:
                current_chunk.append(sentence)
                current_size += sentence_size + 1

        # Add final chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append({
                "content": chunk_text,
                "chunk_index": chunk_index,
                "char_count": len(chunk_text),
            })

        return chunks


class RFPAnalyzer:
    """Analyze RFP documents to extract questions and metadata.
    
    Supports English and Portuguese documents.
    """

    # English patterns
    QUESTION_PATTERNS_EN = [
        r'(?i)^\s*(\d+[.\)]\s+)?(please\s+)?describe[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(please\s+)?explain[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(please\s+)?provide[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?what\s+(is|are)[?:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?how\s+(do|does|will|would)[?:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(list|detail|specify)[.:]?\s*',
        r'(?i)^\s*Q[\d.]+[.:]?\s*',
        r'(?i)^\s*Question\s+[\d.]+[.:]?\s*',
    ]
    
    # Portuguese patterns
    QUESTION_PATTERNS_PT = [
        r'(?i)^\s*(\d+[.\)]\s+)?(por\s+favor\s+)?descreva[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(por\s+favor\s+)?explique[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(por\s+favor\s+)?forneça[.:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(qual|quais)\s+(é|são)[?:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?como\s+(você|a\s+sua\s+solução|a\s+empresa)[?:]?\s*',
        r'(?i)^\s*(\d+[.\)]\s+)?(liste|detalhe|especifique)[.:]?\s*',
        r'(?i)^\s*P[\d.]+[.:]?\s*',
        r'(?i)^\s*Pergunta\s+[\d.]+[.:]?\s*',
        r'(?i)^\s*(quais|qual|como|onde|quando|por\s+que|porquê)\s+',
    ]
    
    # Combined patterns
    QUESTION_PATTERNS = QUESTION_PATTERNS_EN + QUESTION_PATTERNS_PT
    
    # Language detection keywords
    PORTUGUESE_MARKERS = [
        'de', 'para', 'com', 'por', 'em', 'um', 'uma', 'os', 'as', 'que', 
        'qual', 'quais', 'como', 'onde', 'quando', 'porque', 'descreva', 
        'explique', 'forneça', 'liste', 'detalhe', 'sua', 'solução', 
        'empresa', 'produto', 'serviço', 'segurança', 'dados', 'informação'
    ]
    
    # Accented characters common in Portuguese
    PORTUGUESE_ACCENTS = ['ã', 'õ', 'ç', 'á', 'é', 'í', 'ó', 'ú', 'â', 'ê', 'ô']

    @classmethod
    def detect_language(cls, text: str) -> str:
        """Detect document language using heuristics.
        
        Returns: 'en', 'pt', or 'unknown'
        """
        text_lower = text.lower()
        
        # Check for Portuguese accented characters
        accent_count = sum(1 for char in text_lower if char in cls.PORTUGUESE_ACCENTS)
        
        # Check for Portuguese common words
        portuguese_word_count = sum(1 for word in cls.PORTUGUESE_MARKERS if f' {word} ' in f' {text_lower} ')
        
        # Calculate score
        portuguese_score = portuguese_word_count + (accent_count * 0.5)
        
        # Check for English indicators
        english_markers = ['the', 'and', 'for', 'with', 'you', 'your', 'solution', 'please', 'describe', 'explain']
        english_word_count = sum(1 for word in english_markers if f' {word} ' in f' {text_lower} ')
        
        logger.debug("Language detection", 
                    portuguese_score=portuguese_score,
                    english_words=english_word_count,
                    accent_count=accent_count)
        
        if portuguese_score >= 3 or accent_count >= 2:
            return "pt"
        elif english_word_count >= 3:
            return "en"
        
        return "unknown"
    
    @classmethod
    def detect_rfp_type(cls, text: str) -> RFPType:
        """Detect the type of RFP document."""
        text_lower = text.lower()

        # English security keywords
        security_keywords_en = [
            "security questionnaire", "infosec", "cybersecurity", "compliance",
            "soc 2", "iso 27001", "gdpr", "hipaa"
        ]
        # Portuguese security keywords
        security_keywords_pt = [
            "questionário de segurança", "segurança da informação", "segurança cibernética",
            "conformidade", "certificação", "iso 27001", "lgpd"
        ]
        
        all_security_keywords = security_keywords_en + security_keywords_pt
        if any(kw in text_lower for kw in all_security_keywords):
            return RFPType.SECURITY_QUESTIONNAIRE

        if "due diligence" in text_lower or "ddq" in text_lower:
            return RFPType.DDQ
        
        if "due diligence" in text_lower or "auditoria" in text_lower:
            return RFPType.DDQ

        if "request for information" in text_lower or "rfi" in text_lower:
            return RFPType.RFI
        
        if "solicitação de informação" in text_lower or "rfi" in text_lower:
            return RFPType.RFI

        if "proposal" in text_lower and "request" in text_lower:
            return RFPType.RFP
        
        if "proposta" in text_lower and ("solicitação" in text_lower or "edital" in text_lower):
            return RFPType.RFP

        return RFPType.UNKNOWN

    @classmethod
    def extract_questions(cls, text: str) -> list[dict[str, Any]]:
        """Extract questions from RFP text."""
        questions = []
        lines = text.split('\n')
        question_counter = 0

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            # Check if line matches question patterns with Unicode support
            is_question = any(re.match(pattern, line, flags=re.UNICODE) for pattern in cls.QUESTION_PATTERNS)

            # Also check for explicit question marks
            if not is_question and '?' in line and len(line) > 20:
                is_question = True

            if is_question and len(line) > 15:
                question_counter += 1
                questions.append({
                    "id": f"Q-{question_counter:03d}",
                    "question_text": line,
                    "category": cls._categorize_question(line),
                    "priority": cls._determine_priority(line),
                })

        return questions

    @classmethod
    def _normalize_for_classification(cls, text: str, language: str) -> str:
        """Normalize text for classification by translating key terms to English."""
        if language != "pt":
            return text.lower()
        
        # Portuguese to English mapping for classification
        pt_to_en = {
            'criptografia': 'encryption',
            'encriptação': 'encryption',
            'dados em repouso': 'data at rest',
            'dados em trânsito': 'data in transit',
            'dados em transito': 'data in transit',
            'autenticação': 'authentication',
            'controle de acesso': 'access control',
            'autorização': 'authorization',
            'disponibilidade': 'availability',
            'backup': 'backup',
            'recuperação de desastre': 'disaster recovery',
            'plano de continuidade': 'business continuity',
            'segurança': 'security',
            'conformidade': 'compliance',
            'certificação': 'certification',
            'auditoria': 'audit',
            'teste de penetração': 'penetration test',
            'infraestrutura': 'infrastructure',
            'arquitetura': 'architecture',
            'integração': 'integration',
            'preço': 'price',
            'custo': 'cost',
            'orçamento': 'budget',
            'cronograma': 'timeline',
            'prazo': 'deadline',
            'entrega': 'delivery',
            'empresa': 'company',
            'histórico': 'history',
            'equipe': 'team',
            'experiência': 'experience',
            'api': 'api',
            'software': 'software',
            'plataforma': 'platform',
            'nuvem': 'cloud',
            'hospedagem': 'hosting',
            'implementação': 'implementation',
            'suporte': 'support',
            'monitoramento': 'monitoring',
            'logs': 'logging',
            'retenção de dados': 'data retention',
            'resposta a incidentes': 'incident response',
        }
        
        text_lower = text.lower()
        normalized = text_lower
        
        for pt_term, en_term in pt_to_en.items():
            normalized = normalized.replace(pt_term, en_term)
        
        return normalized
    
    @classmethod
    def _categorize_question(cls, text: str) -> str:
        """Categorize question by topic. Supports English and Portuguese."""
        # Detect language first
        language = cls.detect_language(text)
        
        # Normalize for classification
        normalized = cls._normalize_for_classification(text, language)

        categories = {
            "technical": [
                "technology", "architecture", "infrastructure", "api", "integration",
                "software", "platform", "deployment", "hosting", "cloud",
                "tecnologia", "arquitetura", "infraestrutura", "integração",
                "software", "plataforma", "nuvem", "hospedagem", "implementação"
            ],
            "security": [
                "security", "encryption", "authentication", "authorization", "compliance",
                "certification", "audit", "penetration test", "access control",
                "segurança", "criptografia", "autenticação", "autorização", "conformidade",
                "certificação", "auditoria", "teste de penetração", "controle de acesso"
            ],
            "pricing": [
                "price", "cost", "budget", "fee", "pricing", "payment", "invoice",
                "preço", "custo", "orçamento", "taxa", "pagamento", "fatura"
            ],
            "timeline": [
                "timeline", "schedule", "milestone", "deadline", "delivery", "implementation",
                "cronograma", "prazo", "marco", "entrega", "implementação"
            ],
            "company_overview": [
                "company", "about us", "history", "experience", "team", "expertise",
                "empresa", "sobre nós", "histórico", "experiência", "equipe"
            ],
        }

        for category, keywords in categories.items():
            if any(kw in normalized for kw in keywords):
                return category

        return "general"

    @classmethod
    def _determine_priority(cls, text: str) -> int:
        """Determine question priority (1-10). Supports English and Portuguese."""
        text_lower = text.lower()

        # High priority indicators (English + Portuguese)
        high_priority = [
            "mandatory", "required", "must", "critical", "essential",
            "obrigatório", "obrigatorio", "necessário", "necessario", 
            "crítico", "critico", "essencial", "exigido"
        ]
        if any(kw in text_lower for kw in high_priority):
            return 8

        # Medium priority (English + Portuguese)
        medium_priority = [
            "should", "recommend", "prefer", "important",
            "deveria", "recomendado", "preferível", "preferivel", "importante"
        ]
        if any(kw in text_lower for kw in medium_priority):
            return 5

        return 3

    @classmethod
    def extract_deadlines(cls, text: str) -> list[dict[str, str]]:
        """Extract deadline dates from text."""
        deadlines = []

        # Date patterns
        date_patterns = [
            r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
            r'(\w+\s+\d{1,2},?\s+\d{4})',
            r'(\d{1,2}\s+\w+\s+\d{4})',
        ]

        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.UNICODE)
            for match in matches:
                date_str = match.group(1)
                context_start = max(0, match.start() - 100)
                context_end = min(len(text), match.end() + 100)
                context = text[context_start:context_end]

                if any(kw in context.lower() for kw in ["due", "deadline", "submission", "close"]):
                    deadlines.append({
                        "date": date_str,
                        "context": context.strip(),
                    })

        return deadlines
